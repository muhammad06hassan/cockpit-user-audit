#!/usr/bin/env python3
"""
Cockpit Users Audit Script - fix for profile-id/name mapping and distribution counts.

Key fixes:
 - normalize_ids more robustly
 - reconcile users whose `profile` field contains the profile NAME (map to _id)
 - log reconciliation stats and unmapped entries for QA
 - Profile ID column is the _id; Profile Name column is name (correct order)
 - User distribution counts now correct (includes 0 counts); UNASSIGNED row preserved
 - email_validation.csv lists only non-@jeeny.me accounts
"""

import argparse
import csv
import json
import logging
import os
import re
import shutil
import sys
from collections import defaultdict
from datetime import datetime
from typing import Any, Dict, List, Optional, Set, Tuple

from docx import Document

# ---------- Config ----------
SNAPSHOT_DIR = "snapshots"
OUTPUT_ROOT = "Audit_Results"
JEENY_DOMAIN = "@jeeny.me"

PERMISSION_MAP = {32: "Special", 16: "Delete", 8: "Edit", 4: "Add", 2: "View", 1: "List"}

# ---------- Logging ----------
def setup_logging(output_dir: str):
    os.makedirs(output_dir, exist_ok=True)
    log_path = os.path.join(output_dir, "audit.log")
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[logging.FileHandler(log_path, encoding="utf-8"), logging.StreamHandler(sys.stdout)],
    )

# ---------- Utilities ----------
def load_json(path: str) -> Any:
    with open(path, "r", encoding="utf-8") as fh:
        return json.load(fh)

def save_json(path: str, data: Any):
    with open(path, "w", encoding="utf-8") as fh:
        json.dump(data, fh, indent=2, ensure_ascii=False)

def ensure_dir(path: str):
    os.makedirs(path, exist_ok=True)

def export_csv(path: str, header: List[str], rows: List[List[Any]]):
    ensure_dir(os.path.dirname(path) or ".")
    with open(path, "w", newline="", encoding="utf-8") as fh:
        writer = csv.writer(fh)
        writer.writerow(header)
        writer.writerows(rows)

# ---------- Normalization helpers ----------
OID_RE = re.compile(r"ObjectId\(['\"]?([0-9a-fA-F]{8,})['\"]?\)")

def normalize_id(raw: Any) -> Optional[str]:
    """Normalize common MongoID representations into a plain string id."""
    if raw is None:
        return None
    # dict with nested $oid / _id / id
    if isinstance(raw, dict):
        # direct patterns
        for key in ("$oid", "oid", "_id", "id"):
            if key in raw and isinstance(raw[key], str):
                return raw[key].strip()
        # nested candidate values - prefer nested dict with $oid or _id
        for v in raw.values():
            if isinstance(v, dict):
                for key in ("$oid", "_id", "id"):
                    if key in v and isinstance(v[key], str):
                        return v[key].strip()
        # if dict contains a string value that looks like hex id, prefer it
        for v in raw.values():
            if isinstance(v, str):
                s = v.strip()
                if re.match(r"^[0-9a-fA-F]{8,}$", s) or re.match(r"^[0-9a-fA-F]{24}$", s):
                    return s
        # fallback: stringified dict
        try:
            return json.dumps(raw, sort_keys=True)
        except Exception:
            return str(raw)
    # string handling
    if isinstance(raw, str):
        s = raw.strip()
        m = OID_RE.match(s)
        if m:
            return m.group(1)
        # plain hex 24-char
        if re.match(r"^[0-9a-fA-F]{24}$", s):
            return s
        # if it's a quoted string, strip
        s = s.strip("'\" ")
        return s
    # numeric
    if isinstance(raw, (int, float)):
        return str(raw)
    # fallback
    try:
        return str(raw)
    except Exception:
        return None

def normalize_permission(value: Any) -> int:
    if value is None:
        return 0
    if isinstance(value, int):
        return value
    if isinstance(value, str):
        try:
            return int(value)
        except Exception:
            digits = re.findall(r"\d+", value)
            return int(digits[0]) if digits else 0
    if isinstance(value, dict):
        for k in ("$numberInt", "$numberLong"):
            if k in value:
                try:
                    return int(value[k])
                except Exception:
                    pass
        for v in value.values():
            if isinstance(v, int):
                return v
            if isinstance(v, str) and v.isdigit():
                return int(v)
            if isinstance(v, dict):
                for vv in v.values():
                    if isinstance(vv, int):
                        return vv
        return 0
    if isinstance(value, list):
        for it in value:
            if isinstance(it, int):
                return it
            if isinstance(it, str) and it.isdigit():
                return int(it)
    return 0

def decode_permissions(value: int) -> List[str]:
    try:
        v = int(value)
    except Exception:
        return []
    return [label for bit, label in sorted(PERMISSION_MAP.items(), reverse=True) if v & bit]

def format_perm_value(value: Any) -> str:
    intval = normalize_permission(value)
    labels = decode_permissions(intval)
    return f"{intval} ({', '.join(labels)})" if labels else f"{intval} (None)"

# ---------- Data extraction helpers ----------
def extract_list(obj: Any, prefer_keys: Tuple[str, ...] = ("profiles", "users", "data", "results", "items")) -> List[Dict]:
    if isinstance(obj, list):
        return obj
    if isinstance(obj, dict):
        for k in prefer_keys:
            if k in obj and isinstance(obj[k], list):
                return obj[k]
        for v in obj.values():
            if isinstance(v, list):
                return v
    return []

# ---------- Normalizing records ----------
def normalize_profiles(raw_profiles: Any) -> List[Dict]:
    items = extract_list(raw_profiles, ("profiles", "data", "results", "items"))
    out = []
    for p in items:
        if not isinstance(p, dict):
            continue
        _id = normalize_id(p.get("_id") if "_id" in p else p.get("id") if "id" in p else None)
        name = p.get("name") or p.get("profile_name") or p.get("title") or "N/A"
        desc = p.get("description") or p.get("desc") or ""
        perms_raw = p.get("permissions", {}) or {}
        perms = {}
        if isinstance(perms_raw, dict):
            for k, v in perms_raw.items():
                perms[k] = normalize_permission(v)
        out.append({"_id": _id, "name": name, "description": desc, "permissions": perms})
    return out

def normalize_users(raw_users: Any) -> List[Dict]:
    items = extract_list(raw_users, ("users", "data", "results", "items"))
    out = []
    for u in items:
        if not isinstance(u, dict):
            continue
        uid = normalize_id(u.get("_id") or u.get("id"))
        # Attempt multiple profile fields
        profile_field = u.get("profile") or u.get("profile_id") or u.get("profileId") or u.get("role") or None
        profile_id = None
        profile_name_from_field = None
        if profile_field is not None:
            if isinstance(profile_field, dict):
                # Prefer nested _id or $oid
                if "_id" in profile_field:
                    profile_id = normalize_id(profile_field["_id"])
                elif "$oid" in profile_field:
                    profile_id = normalize_id(profile_field["$oid"])
                elif "id" in profile_field:
                    profile_id = normalize_id(profile_field["id"])
                # If there is a name inside this object, keep it as possible fallback
                if "name" in profile_field and isinstance(profile_field["name"], str):
                    profile_name_from_field = profile_field["name"].strip()
            else:
                # string or numeric; normalize but may be a name
                profile_id = normalize_id(profile_field)
        # fallback: if profile_id is actually a profile name (we'll reconcile later), keep it as-is
        email = (u.get("email") or "").strip().lower()
        name = u.get("name") or u.get("full_name") or u.get("username") or "N/A"
        status = (u.get("status") or "").strip() or u.get("state") or ""
        record = {"_id": uid or "N/A", "name": name, "email": email, "profile": profile_id, "profile_name_field": profile_name_from_field, "status": status}
        out.append(record)
    return out

# ---------- Reconciliation (map name->id) ----------
def reconcile_user_profile_references(users: List[Dict], profiles: List[Dict]) -> Tuple[int,int,List[str]]:
    """Map users whose profile field contains profile NAME back to the profile _id.
       Returns (mapped_count, unmapped_count, sample_unmapped_list)
    """
    id_set = {p.get("_id") for p in profiles if p.get("_id")}
    name_to_id = { (p.get("name") or "").strip().lower(): p.get("_id") for p in profiles if p.get("_id") }
    mapped = 0
    unmapped = 0
    unmapped_examples = []
    for u in users:
        pid = u.get("profile")
        pname_field = u.get("profile_name_field")
        if pid and pid in id_set:
            continue  # already a valid id
        # if user.profile is None but profile_name_field exists
        if not pid and pname_field:
            lookup = pname_field.strip().lower()
            if lookup in name_to_id:
                u["profile"] = name_to_id[lookup]
                mapped += 1
                continue
        # if pid exists but matches a profile NAME (case-insensitive), map it
        if pid:
            pid_str = str(pid).strip().lower()
            if pid_str in name_to_id:
                u["profile"] = name_to_id[pid_str]
                mapped += 1
                continue
        # final attempt: if pid looks like JSON with name inside, try to find any profile name
        if pid and isinstance(pid, str):
            low = pid.lower()
            for nm, idx in name_to_id.items():
                if nm and nm in low:
                    u["profile"] = idx
                    mapped += 1
                    break
            else:
                unmapped += 1
                if len(unmapped_examples) < 10:
                    unmapped_examples.append(str(pid))
        else:
            # no profile
            if not u.get("profile"):
                # unassigned — treat as unmapped but not an error
                pass
    return mapped, unmapped, unmapped_examples

# ---------- Matrix builders ----------
def build_profile_permissions_matrix(profiles: List[Dict]) -> Tuple[List[str], List[List]]:
    modules = sorted({m for p in profiles for m in p.get("permissions", {}).keys()})
    sorted_profiles = sorted(profiles, key=lambda p: (0 if (p.get("name") or "").strip() == "Super Admin" else 1, (p.get("name") or "").lower()))
    header = ["Profile ID", "Profile Name"] + modules
    rows = []
    for p in sorted_profiles:
        pid = p.get("_id") or "N/A"
        pname = p.get("name") or "N/A"
        row = [pid, pname]
        for module in modules:
            row.append(format_perm_value(p.get("permissions", {}).get(module, 0)))
        rows.append(row)
    return header, rows

def build_user_access_matrix(users: List[Dict], profiles: List[Dict]) -> Tuple[List[str], List[List]]:
    profile_map = {p.get("_id"): p.get("name", "N/A") for p in profiles if p.get("_id")}
    header = ["Profile ID", "Profile Name", "User Name", "User Email", "User ID", "Status"]
    rows = []
    for u in sorted(users, key=lambda x: ((x.get("profile") or ""), (x.get("name") or ""))):
        pid = u.get("profile") or "N/A"
        pname = profile_map.get(pid, "N/A")
        rows.append([pid, pname, u.get("name", "N/A"), u.get("email", "N/A"), u.get("_id", "N/A"), map_status(u)])
    return header, rows

def build_user_distribution(profiles: List[Dict], users: List[Dict]) -> Tuple[List[str], List[List]]:
    counts = defaultdict(int)
    unassigned = 0
    for u in users:
        pid = u.get("profile")
        if pid:
            counts[pid] += 1
        else:
            unassigned += 1
    dist = []
    for p in profiles:
        pid = p.get("_id") or "N/A"
        name = p.get("name") or "N/A"
        dist.append([pid, name, counts.get(pid, 0)])
    if unassigned:
        dist.append(["N/A", "UNASSIGNED", unassigned])
    dist.sort(key=lambda x: (0 if x[1] == "Super Admin" else 1, x[1].lower()))
    numbered = [[i + 1, row[0], row[1], row[2]] for i, row in enumerate(dist)]
    header = ["No.", "Profile ID", "Profile Name", "User Count"]
    return header, numbered

# ---------- Status mapping ----------
def map_status(user: Dict) -> str:
    s = (user.get("status") or "").strip().lower()
    if s == "enabled" or s == "active":
        return "Active"
    if s == "pending":
        return "Pending Approval"
    if s == "disabled":
        return "Inactive"
    return "Inactive"

# ---------- Email validation ----------
def email_validation(users: List[Dict]) -> Tuple[List[str], List[List], List[str], int]:
    header = ["User Email", "User Name", "Status", "Note"]
    rows = []
    active_bad = []
    inactive_count = 0
    for u in users:
        email = (u.get("email") or "").strip().lower()
        if not email:
            continue
        if not email.endswith(JEENY_DOMAIN):
            status = map_status(u)
            note = "Active non-jeeny" if status == "Active" else "Inactive non-jeeny"
            rows.append([email, u.get("name", "N/A"), status, note])
            if status == "Active":
                active_bad.append(email)
            else:
                inactive_count += 1
    return header, rows, active_bad, inactive_count

# ---------- Offboarding validation ----------
def offboarding_validation(users: List[Dict], hr_offboarded: Set[str]) -> Tuple[List[str], List[List], List[str]]:
    header = ["User Email", "User Name", "Cockpit Status", "In HR List", "Validation"]
    rows = []
    discrepancies = []
    for u in users:
        email = (u.get("email") or "").strip().lower()
        status = map_status(u)
        in_hr = email in hr_offboarded
        if in_hr and status != "Inactive":
            note = "❌ Discrepancy: In HR list but not disabled"
            discrepancies.append(f"{email} -> {status}")
        else:
            note = "OK"
        rows.append([email or "N/A", u.get("name", "N/A"), status, str(in_hr), note])
    return header, rows, discrepancies

# ---------- Comparisons ----------
def compare_profiles(old_profiles: List[Dict], new_profiles: List[Dict]) -> Dict:
    old_map = {p["_id"]: p for p in old_profiles if p.get("_id")}
    new_map = {p["_id"]: p for p in new_profiles if p.get("_id")}
    added = sorted(list(set(new_map.keys()) - set(old_map.keys())))
    removed = sorted(list(set(old_map.keys()) - set(new_map.keys())))
    changes = []
    for pid in set(old_map.keys()) & set(new_map.keys()):
        diffs = {}
        o = old_map[pid]; n = new_map[pid]
        if (o.get("name") or "") != (n.get("name") or ""):
            diffs["name"] = {"old": o.get("name"), "new": n.get("name")}
        if (o.get("description") or "") != (n.get("description") or ""):
            diffs["description"] = {"old": o.get("description"), "new": n.get("description")}
        perm_changes = []
        for k in set(o.get("permissions", {}).keys()) | set(n.get("permissions", {}).keys()):
            ov = normalize_permission(o.get("permissions", {}).get(k, 0))
            nv = normalize_permission(n.get("permissions", {}).get(k, 0))
            if ov != nv:
                perm_changes.append({"module": k, "old": format_perm_value(ov), "new": format_perm_value(nv)})
        if perm_changes:
            diffs["permissions"] = perm_changes
        if diffs:
            changes.append({"_id": pid, "changes": diffs})
    return {"added": added, "removed": removed, "changes": changes}

def compare_users(old_users: List[Dict], new_users: List[Dict],
                  old_profiles: List[Dict], new_profiles: List[Dict]) -> Dict:
    """
    Compare users based on _id, detect changes in name, email, status, and profile.
    Profiles are compared by MongoDB _id, but JSON may only store profile names,
    so we normalize profile names → ids using the profile maps.
    """
    old_map = {u["_id"]: u for u in old_users if u.get("_id")}
    new_map = {u["_id"]: u for u in new_users if u.get("_id")}

    # build maps for name→id and id→name
    old_name_to_id = { (p.get("name") or "").strip(): p["_id"] for p in old_profiles if p.get("_id")}
    new_name_to_id = { (p.get("name") or "").strip(): p["_id"] for p in new_profiles if p.get("_id")}

    added = sorted(list(set(new_map.keys()) - set(old_map.keys())))
    removed = sorted(list(set(old_map.keys()) - set(new_map.keys())))
    changes = []

    for uid in set(old_map.keys()) & set(new_map.keys()):
        o = old_map[uid]
        n = new_map[uid]
        diffs = {}

        # simple fields
        for k in ("name", "email", "status"):
            if str((o.get(k) or "")).strip() != str((n.get(k) or "")).strip():
                diffs[k] = {"old": o.get(k), "new": n.get(k)}

        # normalize profile to IDs
        old_profile_raw = (o.get("profile") or "").strip()
        new_profile_raw = (n.get("profile") or "").strip()
        old_pid = old_name_to_id.get(old_profile_raw, old_profile_raw)
        new_pid = new_name_to_id.get(new_profile_raw, new_profile_raw)

        if old_pid != new_pid:
            diffs["profile"] = {"old": old_pid, "new": new_pid}

        if diffs:
            changes.append({"_id": uid, "changes": diffs})

    return {"added": added, "removed": removed, "changes": changes}


# ---------- Baseline handling ----------
def load_baseline_normalized() -> Tuple[Optional[List[Dict]], Optional[List[Dict]]]:
    ppath = os.path.join(SNAPSHOT_DIR, "base_profiles.json")
    upath = os.path.join(SNAPSHOT_DIR, "base_internal_users.json")
    if not (os.path.exists(ppath) and os.path.exists(upath)):
        return None, None
    try:
        profiles = load_json(ppath)
        users = load_json(upath)
        return profiles, users
    except Exception as e:
        logging.warning("Failed to load baseline: %s", e)
        return None, None

def update_baseline_normalized(orig_profiles_path: str, orig_users_path: str, normalized_profiles: List[Dict], normalized_users: List[Dict]):
    ensure_dir(SNAPSHOT_DIR)
    save_json(os.path.join(SNAPSHOT_DIR, "base_profiles.json"), normalized_profiles)
    save_json(os.path.join(SNAPSHOT_DIR, "base_internal_users.json"), normalized_users)
    try:
        shutil.copy2(orig_profiles_path, os.path.join(SNAPSHOT_DIR, "raw_profiles_original.json"))
        shutil.copy2(orig_users_path, os.path.join(SNAPSHOT_DIR, "raw_users_original.json"))
    except Exception:
        pass

# ---------- Report generation ----------
def generate_report(output_dir: str,
                    profiles_current,
                    profiles_baseline,
                    users_current,
                    prof_comp: Optional[Dict],
                    user_comp: Optional[Dict],
                    offboard_discrepancies: List[str],
                    email_active_bad: List[str],
                    inactive_non_jeeny_count: int,
                    auditor: str = "N/A",
                    requestor: str = "N/A"):
    doc = Document()
    doc.add_heading("Cockpit User Access Audit Report", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Auditor: {auditor}")
    doc.add_paragraph(f"Requestor: {requestor}")

    # build profile maps
    profile_map_new = {p.get("_id"): p.get("name", "N/A") for p in (profiles_current or []) if p.get("_id")}
    profile_map_old = {p.get("_id"): p.get("name", "N/A") for p in (profiles_baseline or []) if p.get("_id")} if profiles_baseline else {}

    # Profile Changes
    doc.add_heading("Profile Changes", level=2)
    doc.add_paragraph(f"Total profiles: {len(profiles_current or [])}")
    if prof_comp:
        doc.add_paragraph(f"Added profiles: {len(prof_comp.get('added', []))}")
        if prof_comp.get("added"): doc.add_paragraph(", ".join(prof_comp["added"]))
        doc.add_paragraph(f"Removed profiles: {len(prof_comp.get('removed', []))}")
        if prof_comp.get("removed"): doc.add_paragraph(", ".join(prof_comp["removed"]))
        if prof_comp.get("changes"):
            doc.add_paragraph("Profiles with changes:")
            for c in prof_comp["changes"]:
                doc.add_paragraph(f"- {c['_id']}: {c['changes']}")
        else:
            doc.add_paragraph("No profile changes detected.")
    else:
        doc.add_paragraph("No baseline found to compare profiles.")

    # User Changes
    doc.add_heading("User Changes", level=2)
    doc.add_paragraph(f"Total users: {len(users_current or [])}")
    if user_comp:
        doc.add_paragraph(f"New users: {len(user_comp.get('added', []))}")
        doc.add_paragraph(f"Removed users: {len(user_comp.get('removed', []))}")
        if user_comp.get("changes"):
            doc.add_paragraph("Users with attribute changes:")
            #profile_map_new = {p.get("_id"): p.get("name", "N/A") for p in (profiles_current or []) if p.get("_id")}
            #profile_map_old = {p.get("_id"): p.get("name", "N/A") for p in (profiles_baseline or []) if p.get("_id")} if profiles_baseline else {}
            for c in user_comp["changes"]:
                formatted_changes = format_user_change(c["changes"], profile_map_old, profile_map_new)
                doc.add_paragraph(f"User {c['_id']} changes: {formatted_changes}")
        else:
            doc.add_paragraph("No user attribute changes detected.")
    else:
        doc.add_paragraph("No baseline found to compare users.")

    # Email validation
    doc.add_heading("Email Validation (non-jeeny accounts)", level=2)
    if email_active_bad:
        doc.add_paragraph(f"Active non-{JEENY_DOMAIN} users: {len(email_active_bad)}")
        for e in email_active_bad:
            doc.add_paragraph(f"- {e}")
    else:
        doc.add_paragraph("No active non-jeeny domain users found.")
    if inactive_non_jeeny_count:
        doc.add_paragraph(f"Inactive non-jeeny accounts: {inactive_non_jeeny_count}")

    # Offboarding validation
    doc.add_heading("Offboarding Validation", level=2)
    if offboard_discrepancies:
        doc.add_paragraph(f"Discrepancies found: {len(offboard_discrepancies)}")
        for d in offboard_discrepancies:
            doc.add_paragraph(f"- {d}")
    else:
        doc.add_paragraph("No offboarding discrepancies found.")

    out_path = os.path.join(output_dir, "audit_report.docx")
    doc.save(out_path)
    logging.info("Saved report: %s", out_path)
    return out_path


# ---------- User profile mapping in report ----------
def format_user_change(change: Dict, profile_map_old: Dict[str, str], profile_map_new: Dict[str, str]) -> Dict:
    """
    Convert profile ID changes to profile names for readability.
    - old ID is looked up in profile_map_old (baseline) first, then fallback to profile_map_new, then raw id
    - new ID is looked up in profile_map_new (current) first, then fallback to profile_map_old, then raw id
    """
    formatted = {}
    for key, val in change.items():
        if key == "profile":
            old_id = val.get("old")
            new_id = val.get("new")

            # old → baseline first
            if old_id is not None:
                old_name = profile_map_old.get(old_id) or profile_map_new.get(old_id) or old_id
            else:
                old_name = old_id

            # new → current first
            if new_id is not None:
                new_name = profile_map_new.get(new_id) or profile_map_old.get(new_id) or new_id
            else:
                new_name = new_id

            formatted[key] = {"old": old_name, "new": new_name}
        else:
            formatted[key] = val
    return formatted


# ---------- Main ----------
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--profiles", default="profiles.json", help="Profiles JSON file path")
    parser.add_argument("--users", default="internal_users.json", help="Internal users JSON file path")
    parser.add_argument("--hr", default="offboarded_users.csv", help="HR offboarded CSV (emails)")
    parser.add_argument("--auditor", default="Muhammad Hassan")
    parser.add_argument("--requestor", default="Mudasar Yasin")
    args = parser.parse_args()

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    output_dir = os.path.join(OUTPUT_ROOT, f"Audit_{timestamp}")
    ensure_dir(output_dir)
    setup_logging(output_dir)
    logging.info("Starting Cockpit Audit run")

    # Load input files
    try:
        raw_profiles = load_json(args.profiles)
    except Exception as e:
        logging.error("Failed to load profiles JSON: %s", e)
        sys.exit(1)
    try:
        raw_users = load_json(args.users)
    except Exception as e:
        logging.error("Failed to load users JSON: %s", e)
        sys.exit(1)

    # Normalize source data
    profiles = normalize_profiles(raw_profiles)
    users = normalize_users(raw_users)
    logging.info("Normalized profiles: %d, users: %d", len(profiles), len(users))

    # Reconcile users that reference profile names instead of IDs
    mapped_count, unmapped_count, unmapped_examples = reconcile_user_profile_references(users, profiles)
    logging.info("Reconciled %d user profile references by name mapping; %d unmapped examples", mapped_count, unmapped_count)
    if unmapped_examples:
        logging.info("Sample unmapped profile references: %s", unmapped_examples[:10])

    # Load HR offboard list
    hr_offboarded: Set[str] = set()
    if os.path.exists(args.hr):
        try:
            with open(args.hr, newline="", encoding="utf-8") as fh:
                for row in csv.reader(fh):
                    if row and "@" in row[0]:
                        hr_offboarded.add(row[0].strip().lower())
            logging.info("Loaded HR offboarded list: %d entries", len(hr_offboarded))
        except Exception as e:
            logging.warning("Failed to read HR offboarded CSV: %s", e)

    # Build CSVs
    p_header, p_rows = build_profile_permissions_matrix(profiles)
    export_csv(os.path.join(output_dir, "profile_permissions_matrix.csv"), p_header, p_rows)

    ua_header, ua_rows = build_user_access_matrix(users, profiles)
    export_csv(os.path.join(output_dir, "user_access_matrix.csv"), ua_header, ua_rows)

    ud_header, ud_rows = build_user_distribution(profiles, users)
    export_csv(os.path.join(output_dir, "user_distribution.csv"), ud_header, ud_rows)

    ev_header, ev_rows, active_bad_emails, inactive_non_jeeny_count = email_validation(users)
    export_csv(os.path.join(output_dir, "email_validation.csv"), ev_header, ev_rows)

    of_header, of_rows, of_discrepancies = offboarding_validation(users, hr_offboarded)
    export_csv(os.path.join(output_dir, "offboarding_validation.csv"), of_header, of_rows)

    logging.info("CSV outputs written to %s", output_dir)

    # Baseline comparison
    old_profiles, old_users = load_baseline_normalized()
    prof_comp = user_comp = None
    if old_profiles is not None and old_users is not None:
        prof_comp = compare_profiles(old_profiles, profiles)
        user_comp = compare_users(old_users, users, old_profiles, profiles)
        logging.info("Baseline comparison completed.")
    else:
        logging.info("No baseline snapshots found; baseline will be created for next run.")

    # Save normalized baseline snapshots
    update_baseline_normalized(args.profiles, args.users, profiles, users)
    logging.info("Baseline snapshots updated (normalized)")

    # Generate report
    report_path = generate_report(output_dir, profiles, old_profiles, users, prof_comp, user_comp, of_discrepancies, active_bad_emails, inactive_non_jeeny_count, auditor=args.auditor, requestor=args.requestor)

    logging.info("Audit run finished. Outputs: %s", output_dir)
    print(f"[SUCCESS] Audit finished. Results in: {output_dir}")
    print(f"Report: {report_path}")

if __name__ == "__main__":
    main()